import shutil
import tempfile
import time
import os
from pathlib import Path
from typing import List, Tuple
from fastapi import UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_unstructured import UnstructuredLoader
from unstructured.partition.pdf import partition_pdf
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
from chromadb import Client
from .embeddings import QwenEmbeddings
from .knowledge_base import knowledge_base

# 设置资源文件夹路径
RESOURCES_DIR = Path(__file__).parent.parent.parent / "resources"
RESOURCES_DIR.mkdir(parents=True, exist_ok=True)

# Tesseract 安装路径（如未加到系统环境变量）
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

client = Client()

def extract_pdf_with_text_and_ocr(tmp_path: str, file_name: str) -> List[Document]:
    documents = []

    # 1. 提取 PDF 文字层
    text_loader = PyPDFLoader(tmp_path)
    text_docs = text_loader.load()
    documents.extend([
        Document(page_content=doc.page_content, metadata={"source": file_name, "layer": "text"})
        for doc in text_docs if doc.page_content.strip()
    ])

    # 2. OCR 图像内容
    ocr_elements = partition_pdf(
        filename=tmp_path,
        strategy="hi_res",
        languages=["chi_sim", "eng"],
        infer_table_structure=True
    )
    documents.extend([
        Document(page_content=str(el), metadata={"source": file_name, "layer": "ocr"})
        for el in ocr_elements if str(el).strip()
    ])

    return documents

def extract_docx_with_ocr(tmp_path: str, file_name: str) -> List[Document]:
    docs = []

    # 1. 提取文本内容
    loader = UnstructuredLoader(tmp_path, mode="elements")
    docs.extend(loader.load())

    # 2. 提取图片并 OCR
    docx = DocxDocument(tmp_path)
    rels = docx.part._rels
    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.target_ref:
            img_part = rel_obj.target_part
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_file:
                img_file.write(img_part.blob)
                img_path = img_file.name
            try:
                img = Image.open(img_path)
                ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                if ocr_text.strip():
                    docs.append(Document(page_content=ocr_text.strip(), metadata={"source": file_name, "layer": "ocr"}))
            finally:
                os.unlink(img_path)
    return docs

def process_files(
        files: List[UploadFile],
        qwen_api_key: str,
        qwen_base_url: str
) -> Tuple[List[Document], QwenEmbeddings, List[str]]:
    documents = []
    file_names = []
    temp_files = []
    metadata = []

    collection_name = "my_collection"
    if collection_name not in client.list_collections():
        collection = client.create_collection(collection_name)
    else:
        collection = client.get_collection(collection_name)

    try:
        for file in files:
            safe_name = "".join(c for c in file.filename if c.isalnum() or c in (' ', '.', '_')).rstrip()
            save_path = RESOURCES_DIR / f"{int(time.time())}_{safe_name}"
            with open(save_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            file.file.seek(0)
            file_content = file.file.read()
            with tempfile.NamedTemporaryFile(mode="wb", delete=False, suffix=f"_{file.filename}") as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
                temp_files.append(tmp_path)

            file_extension = os.path.splitext(file.filename)[1].lower()
            if file_extension == ".pdf" or file.content_type == "application/pdf":
                loaded_docs = extract_pdf_with_text_and_ocr(tmp_path, file.filename)
            elif file_extension == ".txt" or file.content_type == "text/plain":
                loader = TextLoader(tmp_path)
                loaded_docs = loader.load()
            elif file_extension == ".docx":
                loaded_docs = extract_docx_with_ocr(tmp_path, file.filename)
            else:
                print(f"跳过不支持的类型: 文件名[{file.filename}] 扩展名[{file_extension}] 服务端类型[{file.content_type}]")
                continue

            for doc in loaded_docs:
                doc.metadata.update({
                    "source": file.filename,
                    "file_type": file.content_type,
                    "upload_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                    "file_size": f"{len(file_content) / 1024:.1f}KB"
                })
                for key, value in doc.metadata.items():
                    if value is None:
                        doc.metadata[key] = ""
                    elif isinstance(value, list):
                        doc.metadata[key] = ', '.join(map(str, value))

                metadata.append({
                    "file_name": file.filename,
                    "file_type": file.content_type,
                    "upload_time": doc.metadata["upload_time"],
                    "file_size": doc.metadata["file_size"]
                })

            for doc in loaded_docs:
                doc.page_content = doc.page_content.replace(" ", "")

            documents.extend(loaded_docs)
            file_names.append(file.filename)

        if not documents:
            raise ValueError("所有文件均无可提取内容")

        MAX_CHUNK_LENGTH1 = 2048  # 最大字符数限制（可类比 token）
        MAX_CHUNK_LENGTH2 = 3072
        buffer1 = ""
        buffer2 = ""
        merged_chunks = []
        for doc in documents:
            content = doc.page_content
            metadata = doc.metadata

            # 使用 RecursiveCharacterTextSplitter 按句子初切，便于合并
            initial_chunks = RecursiveCharacterTextSplitter(
                chunk_size=500, chunk_overlap=0, length_function=len, add_start_index=True
            ).split_text(content)

            for chunk in initial_chunks:
                if len(buffer1) + len(chunk) <= MAX_CHUNK_LENGTH1:
                    buffer1 += chunk
                else:
                    merged_chunks.append(Document(page_content=buffer1.strip(), metadata=metadata.copy()))
                    buffer1 = chunk
                if len(buffer2) + len(chunk) <= MAX_CHUNK_LENGTH2:
                    buffer2 += chunk
                else:
                    merged_chunks.append(Document(page_content=buffer2.strip(), metadata=metadata.copy()))
                    buffer2 = chunk
        if buffer1:
            merged_chunks.append(Document(page_content=buffer1.strip(), metadata=metadata.copy()))
        if buffer2:
            merged_chunks.append(Document(page_content=buffer2.strip(), metadata=metadata.copy()))

        chunks = merged_chunks

        if not chunks:
            raise ValueError("文本分割后无有效内容")

        embeddings = QwenEmbeddings(api_key=qwen_api_key, base_url=qwen_base_url)

        ids = [str(int(time.time() * 1000)) + str(i) for i in range(len(chunks))]
        print(f"Generated IDs: {ids}")

        knowledge_base.update_knowledge_base(chunks, embeddings, file_names, ids)

        return chunks, embeddings, file_names

    finally:
        for tmp_path in temp_files:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
